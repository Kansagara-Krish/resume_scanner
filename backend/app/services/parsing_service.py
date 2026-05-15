import io
import re
from typing import Optional

from fastapi.concurrency import run_in_threadpool


class ParsingService:
    def __init__(self) -> None:
        pass

    def extract_text_from_bytes(self, content: bytes, filename: str) -> str:
        ext = filename.split('.')[-1].lower()
        text = ""

        try:
            if ext == 'pdf':
                import fitz  # PyMuPDF
                doc = fitz.open(stream=content, filetype='pdf')
                for page in doc:
                    text += page.get_text()
                doc.close()
            elif ext in ['doc', 'docx']:
                try:
                    import docx
                    doc = docx.Document(io.BytesIO(content))
                    paragraphs = [p.text for p in doc.paragraphs]
                    tables = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                tables.append(' | '.join(row_text))
                    text = '\n'.join(paragraphs + tables)
                except Exception as e:
                    text = ""
            elif ext in ['csv']:
                import pandas as pd
                df = pd.read_csv(io.BytesIO(content))
                text = df.to_string()
            elif ext in ['xlsx', 'xls']:
                import pandas as pd
                df = pd.read_excel(io.BytesIO(content))
                text = df.to_string()
            else:
                text = content.decode('utf-8', errors='ignore')
        except Exception:
            text = ""

        return text.strip()

    async def extract_resume_text(self, content: bytes, filename: str) -> str:
        text = await run_in_threadpool(self.extract_text_from_bytes, content, filename)
        return self.normalize_text(text)

    @staticmethod
    def normalize_text(text: Optional[str]) -> str:
        if not text:
            return ""
        normalized = re.sub(r"\r\n?", "\n", text)
        normalized = re.sub(r"[ \t]+", " ", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()
